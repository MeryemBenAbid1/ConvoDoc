import fitz  # PyMuPDF
import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def is_arabic(text):
    """Check if text contains Arabic characters."""
    return re.search(r'[\u0600-\u06FF]', text) is not None

def set_paragraph_rtl(paragraph):
    """Set RTL (right-to-left) formatting for Arabic paragraphs."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    
    # Set BiDi flag (RTL)
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def fix_arabic_text(text):
    """
    Finds Arabic substrings and reverses them because PyMuPDF 
    often extracts them in visual (reversed) order.
    Keeps English/Numbers as is.
    """
    if not text:
        return text

    # Regex to find Arabic sequences (including spaces between Arabic words)
    # This prevents reversing English text or numbers
    arabic_pattern = re.compile(r'([\u0600-\u06FF\s]+)')
    
    def reverse_match(match):
        s = match.group(1)
        # Only reverse if it actually contains Arabic letters
        if any('\u0600' <= c <= '\u06FF' for c in s):
            return s[::-1]
        return s

    # Apply reversal only to Arabic parts
    fixed_text = arabic_pattern.sub(reverse_match, text)
    return fixed_text

def build_docx_pipeline(input_path: str, output_filename: str) -> str:
    """
    Convert PDF to DOCX with proper Arabic text handling.
    """
    os.makedirs("temp", exist_ok=True)
    os.makedirs("temp/images", exist_ok=True)

    output_path = os.path.join("temp", f"{output_filename}.docx")

    pdf = fitz.open(input_path)
    docx_out = Document()

    try:
        for page_idx in range(len(pdf)):
            page = pdf[page_idx]
            
            # Get blocks from page
            blocks = page.get_text("dict")["blocks"]

            # Sort blocks by vertical position (top to bottom)
            blocks = sorted(blocks, key=lambda b: b["bbox"][1])

            for block_idx, block in enumerate(blocks):
                
                # TEXT BLOCK
                if block["type"] == 0:
                    text = ""
                    
                    for line in block["lines"]:
                        for span in line["spans"]:
                            text += span["text"] + " "
                    
                    clean_text = text.strip()
                    
                    if clean_text:
                        p = docx_out.add_paragraph()
                        p.paragraph_format.space_after = Pt(12)
                        
                        if is_arabic(clean_text):
                            set_paragraph_rtl(p)
                            # Fix reversed Arabic letters
                            final_text = fix_arabic_text(clean_text)
                            run = p.add_run(final_text)
                            run.font.rtl = True
                        else:
                            p.add_run(clean_text)

                # IMAGE BLOCK
                elif block["type"] == 1 and "image" in block:
                    img_bytes = block["image"]
                    img_path = os.path.join("temp", "images", f"p{page_idx}_b{block_idx}.png")

                    with open(img_path, "wb") as f:
                        f.write(img_bytes)

                    docx_out.add_picture(img_path, width=Inches(5))
                    docx_out.add_paragraph("")  # spacing after image

            # Page break after each page except last
            if page_idx != len(pdf) - 1:
                docx_out.add_page_break()

        docx_out.save(output_path)

    finally:
        pdf.close()

    return output_path